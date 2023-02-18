from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

test_name = 'CG CONTRACTOR LLC'
test_addr = '585 NORTH 5TH ST 3FL, NEWARK NJ 01707'
test_ein = 'N\A'
test_tot = '6500.00'


def name(s):
    ln_len = 56
    name_len = len(s) + 1
    u_len = ln_len - name_len
    return '_' + s + ('_' * u_len)


def addr(s):
    ln_len = 56
    addr_len = len(s) + 1
    u_len = ln_len - addr_len
    return '_' + s + ('_' * u_len)


def ein(s):
    return '_' + s + '_____'


def a1099(s):
    ln_len = 15
    a1099 = len(s) + 1
    u_len = ln_len - a1099
    return '_' + s + ('_' * u_len)


paragraph_format = style.paragraph_format
paragraph_format.line_spacing = Pt(30)

p = document.add_paragraph()


def sub_block(nm, adr, id, ten99):
    p.add_run('1.Name:');
    p.add_run(name(nm) + '\n').underline = True
    p.add_run('   Address:');
    p.add_run(addr(adr) + '\n').underline = True  # 61
    p.add_run('   EIN/ITIN/SS#:');
    p.add_run(ein(id) + '\n').underline = True
    p.add_run('   1099 Amount Paid: $');
    p.add_run(a1099(ten99)).underline = True;
    p.add_run('  W2 Amount Paid: $________________' + '\n\n')


sub_block(test_name, test_addr, test_ein, test_tot)

document.save('test.docx')
