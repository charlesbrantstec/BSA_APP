import decimal
from decimal import Decimal
import pandas as pd
import re
import sys
import Subs
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


document = Document()

# document.add_heading('INCOME TAX RETURN', level=1)
# document.add_heading('2022', level=1)
# document.add_heading('REIFAST CONSTRUCTION INC', level=1)

# paragraph = document.add_paragraph('INCOME TAX RETURN \n2022\n \nREIFAST CONSTRUCTION INC')
# paragraph_format = paragraph.paragraph_format

# paragraph_format.alignment
# # indicating alignment is inherited from the style hierarchy
# paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
# paragraph_format.alignment

# run = paragraph.add_run()
# font = run.font
# font.name = 'Arial Black'
# font.size = Pt(26)


font = document.styles['Normal'].font
font.name

font.name = 'Arial'
font.name

# document = Document()
run = document.add_paragraph("YOOOOO").add_run()
font = run.font

font.name = 'Calibri'
font.size = Pt(26)


document.save('TEST.docx')